from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import os


def generate_markdown(data, output_path):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# 网络安全扫描报告 (fscan)\n\n")
        f.write("## 1. 扫描统计\n")
        f.write(f"- 发现高危风险/弱口令数: {data['stats']['Critical'] + data['stats']['High']}\n")
        f.write(f"- 发现资产信息数: {len(data['web_titles']) + len(data['net_infos'])}\n\n")

        f.write("## 2. 漏洞与弱口令详情\n")
        f.write("| 类型 | 详细内容 |\n| --- | --- |\n")
        for item in data['vulnerabilities']:
            f.write(f"| 风险 | {item.replace('[+]', '').strip()} |\n")

        f.write("\n## 3. Web资产指纹\n")
        for item in data['web_titles']:
            f.write(f"- {item}\n")


def generate_word(data, output_path):
    doc = Document()
    doc.add_heading('网络安全扫描报告', 0)

    # 1. 统计图表 (保存为临时图片并插入)
    labels = data['stats'].keys()
    values = data['stats'].values()
    plt.figure(figsize=(6, 4))
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=140)
    plt.title("Risk Distribution")
    plt.savefig("temp_chart.png")

    doc.add_heading('1. 风险统计图', level=1)
    doc.add_picture("temp_chart.png", width=Inches(4))

    # 2. 漏洞表格
    doc.add_heading('2. 漏洞与风险详情', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '级别'
    hdr_cells[1].text = '详情内容'

    for item in data['vulnerabilities']:
        row_cells = table.add_row().cells
        row_cells[0].text = 'High/Vuln'
        row_cells[1].text = item

    # 3. Web资产
    doc.add_heading('3. 资产指纹信息', level=1)
    for item in data['web_titles']:
        doc.add_paragraph(item, style='List Bullet')

    doc.save(output_path)
    os.remove("temp_chart.png")  # 清理临时文件