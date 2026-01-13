import re
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import os


def parse_fscan_result(file_path):
    results = {
        "vulnerabilities": [],
        "web_titles": [],
        "net_infos": [],
        "stats": {"Critical": 0, "High": 0, "Medium": 0, "Info": 0},
    }

    lines = []
    # 尝试多种编码方案
    encodings = ["utf-8", "gbk", "gb18030", "latin-1"]
    success = False

    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                lines = f.readlines()
            print(f"[*] 成功使用 {enc} 编码解析文件")
            success = True
            break
        except UnicodeDecodeError:
            continue

    if not success:
        # 如果都失败了，强制读取，忽略无法解析的字节
        print("[!] 警告：无法识别标准编码，尝试忽略错误字节读取...")
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()

    # 开始解析行内容
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 匹配规律：[+] 代表风险，[*] 代表资产或信息
        if "[+]" in line:
            results["vulnerabilities"].append(line)
            # 简单的风险等级判定
            lower_line = line.lower()
            if any(x in lower_line for x in ["poc", "vuln", "ms17-010", "rethinkdb"]):
                results["stats"]["Critical"] += 1
            else:
                results["stats"]["High"] += 1
        elif "[*]" in line:
            if "WebTitle" in line:
                results["web_titles"].append(line)
            else:
                results["net_infos"].append(line)
            results["stats"]["Info"] += 1

    return results


def generate_markdown(data, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("# 网络安全扫描报告 (fscan)\n\n")
        f.write("## 1. 扫描统计\n")
        f.write(
            f"- 发现高危风险/弱口令数: {data['stats']['Critical'] + data['stats']['High']}\n"
        )
        f.write(
            f"- 发现资产信息数: {len(data['web_titles']) + len(data['net_infos'])}\n\n"
        )

        f.write("## 2. 漏洞与弱口令详情\n")
        f.write("| 类型 | 详细内容 |\n| --- | --- |\n")
        for item in data["vulnerabilities"]:
            f.write(f"| 风险 | {item.replace('[+]', '').strip()} |\n")

        f.write("\n## 3. Web资产指纹\n")
        for item in data["web_titles"]:
            f.write(f"- {item}\n")


def generate_word(data, output_path):
    doc = Document()
    doc.add_heading("网络安全扫描报告", 0)

    # 1. 统计图表 (保存为临时图片并插入)
    labels = data["stats"].keys()
    values = data["stats"].values()
    plt.figure(figsize=(6, 4))
    plt.pie(values, labels=labels, autopct="%1.1f%%", startangle=140)
    plt.title("Risk Distribution")
    plt.savefig("temp_chart.png")

    doc.add_heading("1. 风险统计图", level=1)
    doc.add_picture("temp_chart.png", width=Inches(4))

    # 2. 漏洞表格
    doc.add_heading("2. 漏洞与风险详情", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "级别"
    hdr_cells[1].text = "详情内容"

    for item in data["vulnerabilities"]:
        row_cells = table.add_row().cells
        row_cells[0].text = "High/Vuln"
        row_cells[1].text = item

    # 3. Web资产
    doc.add_heading("3. 资产指纹信息", level=1)
    for item in data["web_titles"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(output_path)
    os.remove("temp_chart.png")  # 清理临时文件

    import sys


def main():
    input_file = "report.txt"  # 你的 fscan 输出文件
    output_dir = "output"

    # --- 新增：自动创建输出文件夹 ---
    if not os.path.exists(output_dir):
        print(f"[*] 正在创建输出目录: {output_dir}")
        os.makedirs(output_dir)
    # ----------------------------

    print(f"[*] 正在解析 fscan 结果: {input_file}")
    data = parse_fscan_result(input_file)

    print("[*] 正在生成 Markdown 报告...")
    generate_markdown(data, os.path.join(output_dir, "report.md"))

    print("[*] 正在生成 Word 报告...")
    generate_word(data, os.path.join(output_dir, "report.docx"))

    print(f"[+] 报告生成成功！存放于 {output_dir} 目录下。")


if __name__ == "__main__":
    main()
